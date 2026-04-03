import json
import os
import csv
import argparse
import pandas as pd
from sentence_transformers import SentenceTransformer, util
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def generate_report():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input_json', type=str, required=True, help='Path to results JSON')
    parser.add_argument('--output_dir', type=str, required=True, help='Directory to save reports')
    parser.add_argument('--model_name', type=str, default='all-mpnet-base-v2', help='Embedding model')
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)
    basename = os.path.basename(args.input_json).replace('.json', '')
    csv_path = os.path.join(args.output_dir, f"{basename}_full_report.csv")
    docx_path = os.path.join(args.output_dir, f"{basename}_stories_review.docx")

    print(f"Loading data from {args.input_json}")
    with open(args.input_json, 'r') as f:
        data = json.load(f)

    print(f"Loading embedding model ({args.model_name})...")
    model = SentenceTransformer(args.model_name)

    records = []
    
    # Initialize DOCX if available
    if HAS_DOCX:
        doc = Document()
        doc.add_heading(f'WISE Generation Report: {basename}', 0)
    else:
        print("Warning: python-docx not installed. Skipping DOCX generation. Please run: pip install python-docx")

    print(f"Processing {len(data)} stories...")
    for i, item in enumerate(data):
        prompt = item.get('requested_rewrite', {}).get('prompt', '')
        target = item.get('requested_rewrite', {}).get('target_new', '')
        subject = item.get('requested_rewrite', {}).get('subject', '')
        
        post = item.get('post', {})
        # Extract generation (Flexible for different script variants)
        gen_list = post.get('fluency', {}).get('generated_text', [""])
        if not gen_list or gen_list == [""]:
            gen_list = post.get('rewrite_gen_content', [""])
        gen_text = gen_list[0] if gen_list else ""
        
        # Calculate Embedding Score
        if gen_text and gen_text != "<No Generation>":
            emb1 = model.encode(target, convert_to_tensor=True)
            emb2 = model.encode(gen_text, convert_to_tensor=True)
            emb_score = util.cos_sim(emb1, emb2).item()
        else:
            emb_score = 0.0

        # Teacher Accuracy
        acc_list = post.get('rewrite_acc', [0.0])
        acc = acc_list[0] if isinstance(acc_list, list) and len(acc_list) > 0 else 0.0

        record = {
            'ID': i,
            'Subject': subject,
            'Prompt': prompt,
            'Target': target,
            'Generated': gen_text,
            'Teacher Acc': round(acc, 4),
            'Embedding Sim': round(emb_score, 4)
        }
        records.append(record)

        # Add to DOCX
        if HAS_DOCX:
            doc.add_heading(f'Edit {i}: {subject}', level=2)
            p = doc.add_paragraph()
            p.add_run('Prompt: ').bold = True
            p.add_run(prompt)
            
            p = doc.add_paragraph()
            p.add_run('Target: ').bold = True
            p.add_run(target)
            
            p = doc.add_paragraph()
            p.add_run('Generated Story: ').bold = True
            p.add_run(gen_text)
            
            p = doc.add_paragraph()
            p.add_run(f'Teacher Acc: {acc:.4f} | Embedding Sim: {emb_score:.4f}').italic = True
            doc.add_page_break()

    # Save CSV
    df = pd.DataFrame(records)
    df.to_csv(csv_path, index=False)
    print(f"CSV Report saved to {csv_path}")

    # Save DOCX
    if HAS_DOCX:
        doc.save(docx_path)
        print(f"DOCX Report saved to {docx_path}")

if __name__ == "__main__":
    generate_report()
